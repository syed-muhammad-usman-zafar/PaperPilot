
import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import streamlit as st

def create_paper_docx(paper_data, research_query):
   
    try:
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
        # Add document title with better formatting
        title_text = paper_data.get('title', research_query)
        if title_text:
            # Clean and format the title
            clean_title = title_text.strip().title()  # Convert to Title Case
            title = doc.add_heading(clean_title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Make title bold and larger
            for run in title.runs:
                run.font.size = Pt(18)
                run.font.bold = True
        else:
            title = doc.add_heading(research_query.title(), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.size = Pt(18)
                run.font.bold = True
        
        doc.add_paragraph()  # Add space after title
        # Improved section mapping with proper formatting
        section_mapping = {
            "Abstract": "Abstract",
            "Introduction": "Introduction", 
            "Literature Review": "Literature Review",
            "Methodology": "Methodology",
            "Experiments / Results": "Results", # Also try "Experiments / Results"
            "Results": "Results",
            "Experiments": "Experiments",
            "Conclusion": "Conclusion"
        }
        
        # Get sections data and normalize keys
        sections_data = paper_data.get('sections', {})
        
        # Create a normalized mapping of section keys (case-insensitive)
        normalized_sections = {}
        for key, content in sections_data.items():
            normalized_key = key.strip().title()  # Convert to Title Case
            normalized_sections[normalized_key] = content
        
        for section_title, section_key in section_mapping.items():
            # Add properly formatted heading with Title Case and styling
            heading = doc.add_heading(section_title, level=1)
            
            # Style the heading
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            
            # Try multiple variations to find the content
            section_content = None
            search_keys = [section_key, section_key.lower(), section_key.upper(), 
                          section_key.title(), section_title, section_title.lower()]
            
            for search_key in search_keys:
                if search_key in sections_data:
                    section_content = sections_data[search_key]
                    break
                elif search_key in normalized_sections:
                    section_content = normalized_sections[search_key]
                    break
            
            if section_content and len(section_content) > 0:
                for paragraph_text in section_content:
                    if paragraph_text and paragraph_text.strip():
                        clean_text = paragraph_text.strip()
                        if not clean_text.startswith('[') or not clean_text.endswith(']'):
                            p = doc.add_paragraph(clean_text)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            # Improve paragraph formatting
                            for run in p.runs:
                                run.font.size = Pt(12)
            else:
                p = doc.add_paragraph(f"[{section_title} content to be added]")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.italic = True
            # ...existing code...
            
            # Add space between sections
            doc.add_paragraph()
        
        # Fallback: if sections are empty, try to parse raw output
        if not any(sections_data.values()) and 'raw_output' in paper_data:
            raw_content = paper_data['raw_output']
            if raw_content and len(raw_content.strip()) > 100:
                print("[DEBUG] Sections empty, parsing raw output...")
                # Parse raw content by sections
                import re
                
                # Split by common section headers
                section_pattern = r'\*\*(Abstract|Introduction|Literature Review|Methodology|Results|Experiments|Conclusion)\*\*'
                sections = re.split(section_pattern, raw_content, flags=re.IGNORECASE)
                
                if len(sections) > 1:
                    for i in range(1, len(sections), 2):  # Skip the first empty part, then take pairs
                        if i < len(sections) and i+1 < len(sections):
                            section_name = sections[i].strip().title()
                            section_text = sections[i+1].strip()
                            
                            if section_text and len(section_text) > 20:
                                heading = doc.add_heading(section_name, level=1)
                                for run in heading.runs:
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                                
                                # Split into paragraphs
                                paragraphs = [p.strip() for p in section_text.split('\n\n') if p.strip()]
                                for para in paragraphs:
                                    if para and not para.startswith('['):
                                        p = doc.add_paragraph(para)
                                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                        for run in p.runs:
                                            run.font.size = Pt(12)
                                
                                doc.add_paragraph()  # Space between sections
        
        # Add References section - check both structured and raw output
        references_added = False
        
        # First, try to use structured references if available
        if 'references' in paper_data and paper_data['references']:
            ref_heading = doc.add_heading('References', level=1)
            # Style the References heading
            for run in ref_heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                
            refs_content = paper_data['references']
            
            # Clean up references formatting
            if refs_content.startswith('References\n'):
                refs_content = refs_content.replace('References\n', '', 1)
            
            ref_lines = [line.strip() for line in refs_content.split('\n') if line.strip()]
            valid_refs = []
            
            for ref_line in ref_lines:
                # Filter out placeholder text and incomplete references
                if (ref_line and 
                    not ref_line.startswith('[Add ') and 
                    'more relevant references' not in ref_line and
                    not ref_line.startswith('[P') and 
                    len(ref_line) > 20):  # Ensure it's not just a fragment
                    p = doc.add_paragraph(ref_line)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Style reference text
                    for run in p.runs:
                        run.font.size = Pt(11)
                    valid_refs.append(ref_line)
            
            references_added = True
            
        # If no structured references, try to extract from raw output
        elif 'raw_output' in paper_data:
            raw_output = paper_data.get('raw_output', '')
            print(f"[DEBUG] Checking raw_output for references (length: {len(raw_output)})")
            
            if 'References' in raw_output or 'REFERENCES' in raw_output:
                import re
                ref_match = re.search(r'(?:References|REFERENCES)\s*\n(.*?)(?:\n\n|\Z)', raw_output, re.DOTALL | re.IGNORECASE)
                if ref_match:
                    doc.add_heading('References', level=1)
                    ref_content = ref_match.group(1).strip()
                    
                    ref_lines = [line.strip() for line in ref_content.split('\n') if line.strip()]
                    valid_refs = []
                    
                    for ref_line in ref_lines:
                        if (ref_line and 
                            not ref_line.startswith('[Add ') and 
                            'more relevant references' not in ref_line and
                            not ref_line.startswith('[P') and
                            len(ref_line) > 20):
                            p = doc.add_paragraph(ref_line)
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            valid_refs.append(ref_line)
                    
                    references_added = True
        
        if not references_added:
            # Add a placeholder if no references found
            doc.add_heading('References', level=1)
            p = doc.add_paragraph("References will be added here.")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Save to bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error creating DOCX: {str(e)}")
        return None

def generate_filename(research_query):
    import re
    
    # Clean the research query for filename
    clean_name = re.sub(r'[^\w\s-]', '', research_query)
    clean_name = re.sub(r'[-\s]+', '_', clean_name)
    clean_name = clean_name.strip('_')
    
    # Limit length and add extension
    if len(clean_name) > 50:
        clean_name = clean_name[:50]
    
    return f"{clean_name}_research_paper.docx"

def create_download_button(paper_data, research_query, key="download_docx"):
    try:
        with st.spinner("📄 Generating editable DOCX file..."):
            docx_bytes = create_paper_docx(paper_data, research_query)
        
        if docx_bytes:
            filename = generate_filename(research_query)
            st.download_button(
                label="📥 Download as Editable DOCX",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=key,
                help="Download the research paper as an editable Microsoft Word document"
            )
            return True
        else:
            st.error("Failed to generate DOCX file")
            return False
            
    except Exception as e:
        st.error(f"Error creating download: {str(e)}")
        return False

def preview_docx_content(paper_data, research_query):
    st.markdown("### 📋 DOCX Export Preview:")
    sections_data = paper_data.get('sections', {})
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Document Structure:**")
        st.write(f"📄 Title: {research_query}")
        
        section_count = 0
        for section_name, content in sections_data.items():
            if content and len(content) > 0:
                real_content = [p for p in content if p and not p.strip().startswith('[')]
                if real_content:
                    section_count += 1
                    word_count = sum(len(p.split()) for p in real_content)
                    st.write(f"📝 {section_name}: {len(real_content)} paragraphs (~{word_count} words)")
        
        st.write(f"✅ Total sections with content: {section_count}")
    
    with col2:
        st.markdown("**Features:**")
        st.write("✅ Formatted headings and sections")
        st.write("✅ Justified paragraph alignment") 
        st.write("✅ Proper margins and spacing")
        st.write("✅ References section")
        st.write("✅ Fully editable in Microsoft Word")
        total_chars = sum(len(str(content)) for content in sections_data.values())
        estimated_size_kb = max(20, total_chars // 50)  # Rough estimate
        st.write(f"📊 Estimated size: ~{estimated_size_kb} KB")
